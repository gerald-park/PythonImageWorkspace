from docx import Document
import pandas as pd
import tkinter.ttk as ttk
import tkinter.messagebox as msgbox
from tkinter import * # __all__
from tkinter import filedialog
import threading, time

OP_ROW = 1 
HEADER_ROW = 2 
FIRST_ROW = 3 

def FindOpCols(table): 
    first_col = 0
    second_col = 0
    first=''
    for idx, _ in enumerate(table.columns):
        txt = table.cell(OP_ROW, idx).text.lower()
        if 'op' in txt:
            if not first:
                first = txt
                first_col = idx
            if first not in txt:
                second_col = idx
                break
    if second_col == 0:
        second_col = first_col
    return first_col, second_col

def FindRefCols(table): 
    first_col = 0
    second_col = 0
    pre_idx = 0
    first = ''
    for idx, _ in enumerate(table.columns):
        txt = table.cell(HEADER_ROW, idx).text.lower()        
        if 'ref' in txt:            
            if not first:
                first_col = idx
                first = txt
                print('first',idx , txt)                
            if idx > (pre_idx+2) : # ref인데 IDX값이 두칸이상 차이 나면 
                print('second',idx , txt)
                second_col = idx
                break
            pre_idx = idx
    return first_col, second_col

def FineItemCols(table): 
    first_col = 0
    second_col = 0
    pre_idx = 0
    first = ''
    for idx, _ in enumerate(table.columns):
        txt = table.cell(HEADER_ROW, idx).text.lower()        
        if 'man.item no' in txt:            
            if not first:
                first_col = idx
                first = txt
                pre_idx = idx
                print('first',idx , txt)                
            if idx > (pre_idx+2) : # ref인데 IDX값이 두칸이상 차이 나면 
                print('second',idx , txt)
                second_col = idx
                break
            pre_idx = idx
    return first_col, second_col

def FineDescriptionCols(table): 
    first_col = 0
    second_col = 0
    pre_idx = 0
    first = ''
    for idx, _ in enumerate(table.columns):
        txt = table.cell(HEADER_ROW, idx).text.lower()
        if 'description' in txt:
            if not first:
                first_col = idx
                first = txt
                pre_idx = idx     
            if idx > (pre_idx+2) : # ref인데 IDX값이 두칸이상 차이 나면 
                second_col = idx
                break
            pre_idx = idx
    return first_col, second_col

def FineQtyCols(table): 
    first_col = 0
    second_col = 0
    pre_idx = 0
    first = ''
    for idx, _ in enumerate(table.columns):
        txt = table.cell(HEADER_ROW, idx).text.lower()
        if 'qty' in txt:
            if not first:
                first_col = idx
                first = txt
                pre_idx = idx     
            if idx > (pre_idx+2) : # ref인데 IDX값이 두칸이상 차이 나면 
                second_col = idx
                break
            pre_idx = idx
    return first_col, second_col

def getLeftHeaderCols(tb):
    op_col, _ = FindOpCols(tb)
    ref_col, _ = FindRefCols(tb)
    item_col, _ = FineItemCols(tb)
    des_col, _ = FineDescriptionCols(tb)
    qty_col, _ = FineQtyCols(tb) 
    return {'op':op_col ,'ref':ref_col,'item':item_col,'des':des_col,'qty':qty_col}
    

def getRightHeaderCols(tb):
    _, op_col = FindOpCols(tb)
    _, ref_col = FindRefCols(tb)
    _, item_col = FineItemCols(tb)
    _, des_col = FineDescriptionCols(tb)
    _, qty_col = FineQtyCols(tb)     
    return {'op':op_col ,'ref':ref_col,'item':item_col,'des':des_col,'qty':qty_col}

def GetLeftItemsList(table):
    item = []
    items = []
    headers = getLeftHeaderCols(table)
    for i, _ in enumerate(table.rows) : # left
        try:
            op = table.cell(OP_ROW, headers['op']).text
            ref = table.cell(FIRST_ROW + i, headers['ref']).text
            item = table.cell(FIRST_ROW + i, headers['item']).text
            des = table.cell(FIRST_ROW  + i, headers['des']).text
            qty = table.cell(FIRST_ROW + i, headers['qty']).text
        except:
            continue
        if qty.isnumeric():
            item = [op, ref, item, des, qty]
            items.append(item)
    return items

def GetRightItemsList(table):
    item = []
    items = []
    headers = getRightHeaderCols(table)
    for i, _ in enumerate(table.rows) : # right
        
        try:
            op = table.cell(OP_ROW, headers['op']).text
            ref = table.cell(FIRST_ROW + i, headers['ref']).text.replace('\n','')
            item = table.cell(FIRST_ROW + i, headers['item']).text
            des = table.cell(FIRST_ROW  + i, headers['des']).text
            qty = table.cell(FIRST_ROW + i, headers['qty']).text
        except:
            continue
        if qty.isnumeric():
            item = [op, ref, item, des, qty]
            items.append(item)
    return items

# GUI Functions
sopfile = ''

def open_file():
    global sopfile
    global file_name
    try:
        sopfile = filedialog.askopenfilename(initialdir=r'C:\Users\ParkGY\DocumentsCFLTYanadoo\DT Academy\SOPMeterial',title="select a file",
                                            filetypes =(("Word File","*.docx"),
                                            ("all files","*.*")))
    except:
        pass
    
def convert():
    t = threading.Thread(target=threadfunc)
    t.start()

def threadfunc():
    doc = Document(sopfile)
    df = pd.DataFrame(columns=['OperationStep','Ref.','Man.Item.No','Description','Qty'])
    
    for i, tb in enumerate(doc.tables):
        print(i, len(tb.rows), len(tb.columns))
        if len(tb.rows) > 10 : # 최소 열줄은 넘겨야...
            lefts = GetLeftItemsList(tb)
            rights = GetRightItemsList(tb)
            for left in lefts:
                df.loc[len(df)] = left
            for right in rights:    
                df.loc[len(df)] = right
        progress = (i + 1) / len(doc.tables) * 100 # 실제 percent 정보를 계산
        p_var.set(progress)
        progress_bar.update()
        
    f_name = sopfile.split('.')[0] + '.csv' 
    df.to_csv(f_name, encoding='utf-8-sig', index=False, mode='w', header=True)      


root = Tk()
root.title("Docx SOP Converter")
root.geometry("300x150")

# Frame 
frame = Frame(root)
frame.pack(fill="x", padx=5, pady=5) # 간격 띄우기

btn_add_file = Button(frame, padx=5, pady=5, text="Open SOP", command=open_file)
btn_add_file.pack(fill="x",side="top", padx=5, pady=5)

btn_start = Button(frame, padx=5, pady=5, text="Convert", command=convert)
btn_start.pack(fill="x",side="top", padx=5, pady=5)

# 프로그래스 바를 추가
p_var = DoubleVar()
progress_bar = ttk.Progressbar(frame, maximum=100, variable=p_var)
progress_bar.pack(fill="x", padx=5, pady=5)
# Runs
root.mainloop()