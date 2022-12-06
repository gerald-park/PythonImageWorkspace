from docx import Document
import pandas as pd
import tkinter.ttk as ttk
import tkinter.messagebox as msgbox
from tkinter import * # __all__
from tkinter import filedialog
import threading, time

OP_ROW = 1 
HEADER_ROW = 2 
FIRST_ROW = 3 

def getLeftHeaderCols(tb):
    op_cols = []
    ref_cols = []
    item_cols = []
    des_cols = []
    qty_cols = []
    for idx, _ in enumerate(tb.columns):
        if 'op' in tb.cell(OP_ROW, idx).text.lower().strip():
            op_cols.append(idx)
        if 'ref' in tb.cell(HEADER_ROW, idx).text.lower().strip():
            ref_cols.append(idx)
        if 'man.item' in tb.cell(HEADER_ROW, idx).text.lower().strip():
            item_cols.append(idx)
        if 'description' in tb.cell(HEADER_ROW, idx).text.lower().strip():
            des_cols.append(idx)
        if 'qty' in tb.cell(HEADER_ROW, idx).text.lower().strip():
            qty_cols.append(idx)
    if ref_cols:
        return {'ref':ref_cols[0],'item':item_cols[0],'des':des_cols[0],'qty':qty_cols[0]}

def getRightHeaderCols(tb):
    op_cols = []
    ref_cols = []
    item_cols = []
    des_cols = []
    qty_cols = []
    for idx, _ in enumerate(tb.columns):
        if 'op' in tb.cell(OP_ROW, idx).text.lower().strip():
            op_cols.append(idx)
        if 'ref' in tb.cell(HEADER_ROW, idx).text.lower().strip():
            ref_cols.append(idx)
        if 'man.item' in tb.cell(HEADER_ROW, idx).text.lower().strip():
            item_cols.append(idx)
        if 'description' in tb.cell(HEADER_ROW, idx).text.lower().strip():
            des_cols.append(idx)
        if 'qty' in tb.cell(HEADER_ROW, idx).text.lower().strip():
            qty_cols.append(idx)
            
    right_header_cols = {'ref': int ,'item':int,'des':int,'qty':int}
    try:
        right_header_cols['ref'] = ref_cols[2]
        for col in item_cols:
            if col > right_header_cols['ref']:
                right_header_cols['item'] = col
                break
        for col in des_cols:
            if col > right_header_cols['ref']:
                right_header_cols['des'] = col
                break
        for col in qty_cols:
            if col > right_header_cols['ref']:
                right_header_cols['qty'] = col
                break
    except:
        pass
    return right_header_cols

def GetLeftItemsList(table):
    item = []
    items = []
    headers = getLeftHeaderCols(table)
    for i, cells in enumerate(table.rows) : # left
        op = table.cell(OP_ROW, 3).text
        try:
            ref = table.cell(FIRST_ROW + i, headers['ref']).text
            item = table.cell(FIRST_ROW + i, headers['item']).text
            des = table.cell(FIRST_ROW  + i, headers['des']).text
            qty = table.cell(FIRST_ROW + i, headers['qty']).text
        except:
            continue
        if qty.isnumeric():
            item = [op, ref, item, des, qty]
            items.append(item)
    return items

def GetRightItemsList(table):
    item = []
    items = []
    headers = getRightHeaderCols(table)
    for i,cells in enumerate(table.rows) : # left
        op = table.cell(OP_ROW, 19).text
        try:
            ref = table.cell(FIRST_ROW + i, headers['ref']).text.replace('\n','')
            item = table.cell(FIRST_ROW + i, headers['item']).text
            des = table.cell(FIRST_ROW  + i, headers['des']).text
            qty = table.cell(FIRST_ROW + i, headers['qty']).text
        except:
            continue
        if qty.isnumeric():
            item = [op, ref, item, des, qty]
            items.append(item)
    return items

    sopfile = ''

def add_file():
    global sopfile
    global file_name
    try:
        sopfile = filedialog.askopenfilename(initialdir=r'C:\Users\ParkGY\DocumentsCFLTYanadoo\DT Academy\SOPMeterial',title="select a file",
                                            filetypes =(("Word File","*.docx"),
                                            ("all files","*.*")))
    except:
        pass
    
def convert():
    t = threading.Thread(target=threadfunc)
    t.start()
    
def threadfunc():
    doc = Document(sopfile)
    df = pd.DataFrame(columns=['OperationStep','Ref.','Man.Item.No','Description','Qty'])
    
    for i, tb in enumerate(doc.tables):
        print(i, len(tb.rows), len(tb.columns))
        if len(tb.rows) > 10 : # 최소 열줄은 넘겨야...
            lefts = GetLeftItemsList(tb)
            rights = GetRightItemsList(tb)
            for left in lefts:
                df.loc[len(df)] = left
            for right in rights:    
                df.loc[len(df)] = right
        # Progress bar update 
        progress = (i + 1) / len(doc.tables) * 100 # 실제 percent 정보를 계산
        p_var.set(progress)
        progress_bar.update()
        
    f_name = sopfile.split('.')[0] + '.csv' 
    df.to_csv(f_name, encoding='utf-8-sig', index=False, mode='w', header=True) 

root = Tk()
root.title("Docx SOP Converter")
root.geometry("300x150")

# Frame 
file_frame = Frame(root)
file_frame.pack(fill="x", padx=5, pady=5) # 간격 띄우기

btn_add_file = Button(file_frame, padx=5, pady=5, text="Open SOP", width=20, command=add_file)
btn_add_file.pack(side="top")

btn_start = Button(file_frame, padx=5, pady=5, text="Convert", width=20, command=convert)
btn_start.pack(side="top", padx=5, pady=5)

# 진행 상황 Progress Bar
frame_progress = LabelFrame(root, text="진행상황")
frame_progress.pack(fill="x", padx=5, pady=5, ipady=5)

p_var = DoubleVar()
progress_bar = ttk.Progressbar(frame_progress, maximum=100, variable=p_var)
progress_bar.pack(fill="x", padx=5, pady=5)

# Runs
root.mainloop()